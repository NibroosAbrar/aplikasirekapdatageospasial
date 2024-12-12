from pathlib import Path
import pandas as pd
import geopandas as gpd
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
from PyQt5 import QtCore, QtWidgets
from PyQt5.QtWidgets import QFileDialog, QWidget, QTableView, QTableWidgetItem, QMessageBox, QMainWindow, QApplication, QVBoxLayout, QScrollArea, QLabel, QProgressBar
from PyQt5.uic import loadUi
from PyQt5.QtCore import QUrl, Qt, QAbstractTableModel, QFileSystemWatcher, QTimer
import openpyxl
import sys
from PyQt5.QtWebEngineWidgets import QWebEngineView
import folium
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import io
import matplotlib.pyplot as plt
import os
import leafmap
import ipywidgets as widgets
from shapely.geometry import Point
from pyproj import CRS
from pyproj.exceptions import CRSError
from PyQt5.QtGui import QStandardItemModel, QStandardItem
import re

class PandasModel(QAbstractTableModel):
    """Model untuk menampilkan DataFrame di QTableView."""
    def __init__(self, data):
        super(PandasModel, self).__init__()
        self._data = data

    def rowCount(self, parent=None):
        return self._data.shape[0]

    def columnCount(self, parent=None):
        return self._data.shape[1]

    def data(self, index, role=Qt.DisplayRole):
        if index.isValid() and role == Qt.DisplayRole:
            return str(self._data.iloc[index.row(), index.column()])
        return None

    def headerData(self, section, orientation, role=Qt.DisplayRole):
        if role == Qt.DisplayRole:
            if orientation == Qt.Horizontal:
                return str(self._data.columns[section])
            if orientation == Qt.Vertical:
                return str(self._data.index[section])
        return None
    

class MplCanvas(FigureCanvas):
    def __init__(self, parent=None, width=16, height=8, dpi=300):
        fig = Figure(figsize=(width, height), dpi=dpi)
        self.axes = fig.add_subplot(111)
        super(MplCanvas, self).__init__(fig)

class MainApp(QMainWindow):
    def __init__(self):
        super(MainApp, self).__init__()
        print("Loading UI...")  # Debugging statement
                
        self.base_path = Path(__file__).resolve().parent
        
        # Dynamic path untuk UI file
        ui_path = self.base_path / "TampilanREVISI1.ui"
        if not ui_path.exists():
            raise FileNotFoundError(f"File UI tidak ditemukan: {ui_path}")
        loadUi(str(ui_path), self)
        print("UI Loaded Successfully")  # Debugging statement
        

        # Inisialisasi progress bar
        self.progress_load.setValue(0)  # Atur nilai awal ke 0
        self.progress_load.setVisible(False)

        self.dataframe = None #Untuk menyimpan data yang ditampilkan di QTableWidget

        # Inisialisasi
        self.setWindowFlags(QtCore.Qt.FramelessWindowHint)
        self.is_dragging = False
        self.drag_position = QtCore.QPoint()
        
        # Menentukan halaman default saat aplikasi dibuka

        # Menambahkan peta ke frame_30
        # self.home_button.clicked.connect(self.init_leafmap)
        self.init_leafmap()
        self.stackedWidget.setCurrentWidget(self.tampilan_map)  # Pastikan tampilan_map adalah default
        self.stackedWidget_2.setCurrentWidget(self.tableWidget_rekap)
        
        # Menghubungkan tombol dan dropdown ComboBox
        self.comboBox.currentIndexChanged.connect(self.update_stack_widget)
        
        self.tentukan_reg.editingFinished.connect(self.update_chart)  # Provinsi ComboBox
        self.tentukan_wil.editingFinished.connect(self.update_chart)  # Wilayah ComboBox
        self.tentukan_est.editingFinished.connect(self.update_chart)  # Estate ComboBox
        self.tentukan_div.editingFinished.connect(self.update_chart)  # Divisi ComboBox
        
        self.tentukan_reg.setPlaceholderText("Tentukan Region")  # Provinsi ComboBox
        self.tentukan_wil.setPlaceholderText("Tentukan Wilayah")  # Wilayah ComboBox
        self.tentukan_est.setPlaceholderText("Tentukan Estate")  # Estate ComboBox
        self.tentukan_div.setPlaceholderText("Tentukan Divisi")  # Divisi ComboBox
        
        self.tahun_reg.editingFinished.connect(self.update_chart)
        self.tahun_wil.editingFinished.connect(self.update_chart)
        self.tahun_est.editingFinished.connect(self.update_chart)
        self.tahun_div.editingFinished.connect(self.update_chart)
        
        self.tahun_est.setPlaceholderText("Pilih Tahun dan/atau Rotasi")
        self.tahun_wil.setPlaceholderText("Pilih Tahun dan/atau Rotasi")
        self.tahun_div.setPlaceholderText("Pilih Tahun dan/atau Rotasi")
        self.tahun_reg.setPlaceholderText("Pilih Tahun dan/atau Rotasi")
       
        self.hapus_est.setPlaceholderText("Pilih Chart yang Ingin Dihapus")
        self.hapus_wil.setPlaceholderText("Pilih Chart yang Ingin Dihapus")
        self.hapus_div.setPlaceholderText("Pilih Chart yang Ingin Dihapus")
        self.hapus_reg.setPlaceholderText("Pilih Chart yang Ingin Dihapus")
        
        self.hapus_est.editingFinished.connect(self.update_chart)
        self.hapus_wil.editingFinished.connect(self.update_chart)
        self.hapus_div.editingFinished.connect(self.update_chart)
        self.hapus_reg.editingFinished.connect(self.update_chart)
        
        self.pushButton.clicked.connect(self.masukkan_data)
        self.pushButton_3.clicked.connect(self.update_chart)  # Tombol untuk refresh chart
        self.pushButton_4.clicked.connect(self.simpan_data)
        self.pushButton_2.clicked.connect(self.download)
        self.pushButton_6.clicked.connect(self.minimize_window)
        self.pushButton_7.clicked.connect(self.toggle_fullscreen)
        self.pushButton_8.clicked.connect(self.close_window)
        self.pushButton_5.clicked.connect(self.toggle_menu)
        self.pushButton_10.clicked.connect(self.tentukan_areal_statement)  # Button untuk memilih areal statement
        self.pushButton_9.clicked.connect(self.rekap_data)  # Button untuk rekap data
        self.home_button.clicked.connect(self.menu_awal)
        
            # Variabel untuk menyimpan informasi areal statement
        self.areal_statement = None

        # Sembunyikan panel "Visualisasi Data" saat aplikasi pertama kali dibuka
        self.left_menu_widget.setVisible(False)
        self.is_fullscreen = True

        # Path file Excel khusus untuk Bar Chart
        self.excel_file_path = self.base_path / "04. Data Grafik/Data Buat Nibroos 2.xlsx"
        
        # Watcher untuk memantau perubahan file Excel
        self.file_watcher = QFileSystemWatcher()
        self.file_watcher.addPath(str(self.excel_file_path))
        
        # Sambungkan sinyal perubahan file dengan fungsi reload
        self.file_watcher.fileChanged.connect(self.reload_excel_data)
        
        self.load_chart_data()
        
        self.notification_shown = False
        self.temp_csv_path = None  # Untuk menyimpan path CSV sementara
        self.tableWidget_rekap = self.findChild(QWidget, "tableWidget_rekap")  # Pastikan tableWidget_rekap ada
        self.tableView_rekap = None  # Akan diinisialisasi di runtime
        self.is_analysis_done = False
    
    def tentukan_areal_statement(self):
        """Memilih file areal statement (Excel) dan memuatnya ke dalam aplikasi."""
        try:
            file_path, _ = QFileDialog.getOpenFileName(self, "Pilih File Excel", "", "Excel Files (*.xlsx *.xls)")
            if not file_path:
                return

            # Membaca file Excel
            df = pd.read_excel(file_path)
            if df.empty:
                QMessageBox.warning(self, "Peringatan", "File Excel tidak memiliki data!")
                return

            # Validasi kolom yang dibutuhkan
            required_columns = ['BLOK_SAP', 'LUAS_TANAM']
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                QMessageBox.warning(self, "Error", f"Kolom berikut tidak ditemukan: {', '.join(missing_columns)}")
                return

            # Menyimpan DataFrame ke variabel
            self.areal_statement = df

            # Tampilkan DataFrame di tableWidget_2
            self.display_in_tablewidget_2(df)

            QMessageBox.information(self, "Sukses", "File areal statement berhasil dimuat dan ditampilkan.")
            print(f"Areal statement berhasil dimuat: {self.areal_statement.head()}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal memuat file areal statement: {e}")

    def display_in_tablewidget_2(self, dataframe):
        """Menampilkan DataFrame di tableWidget_2."""
        try:
            self.tableWidget_2.clear()
            self.tableWidget_2.setRowCount(dataframe.shape[0])
            self.tableWidget_2.setColumnCount(dataframe.shape[1])
            self.tableWidget_2.setHorizontalHeaderLabels(dataframe.columns)

            for row_idx, row in dataframe.iterrows():
                for col_idx, value in enumerate(row):
                    item = QTableWidgetItem(str(value))
                    item.setTextAlignment(Qt.AlignCenter)  # Atur teks agar rata tengah
                    self.tableWidget_2.setItem(row_idx, col_idx, item)

            # Sesuaikan ukuran kolom dan baris agar konten terlihat
            self.tableWidget_2.resizeColumnsToContents()
            self.tableWidget_2.resizeRowsToContents()

            print("File areal statement berhasil ditampilkan di tableWidget_2.")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal menampilkan data: {e}")
      
    def menu_awal(self):
        """Tampilkan halaman tampilan_map"""
        self.stackedWidget.setCurrentWidget(self.tampilan_map)
        self.stackedWidget_2.setCurrentWidget(self.tableWidget_rekap)
        
    def init_leafmap(self):
        """Inisialisasi Leafmap dan tampilkan di QWebEngineView"""
        try:
            # Membuat peta Leafmap
            m = leafmap.Map(center=[0, 0], zoom=2)
            m.add_basemap("SATELLITE")

            # # Tambahkan marker sebagai contoh
            # m.add_marker(location=[51.509865, -0.118092], popup="London")
            # m.add_marker(location=[40.712776, -74.005974], popup="New York")

            # Simpan peta sebagai file HTML
            self.map_file = os.path.join(os.getcwd(), "leafmap.html")
            m.to_html(self.map_file)
            print(f"Peta berhasil disimpan di: {self.map_file}")

            # Pastikan file HTML berhasil dibuat
            if not os.path.exists(self.map_file):
                raise FileNotFoundError(f"File HTML tidak ditemukan di {self.map_file}")

            # Tampilkan file HTML di QWebEngineView
            self.map_view = QWebEngineView()
            self.map_view.load(QUrl.fromLocalFile(self.map_file))

            # Tambahkan QWebEngineView ke layout frame_30
            if not self.frame_30.layout():
                layout = QVBoxLayout(self.frame_30)  # Tambahkan layout jika belum ada
                self.frame_30.setLayout(layout)

            # Tambahkan map_view ke layout frame_30
            self.frame_30.layout().addWidget(self.map_view)
            print("Leafmap berhasil ditampilkan di frame_30.")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal memuat peta: {e}")
            print(f"Error: {e}")

    def load_chart_data(self):
        """Membaca data dari file Excel untuk bar chart dan mengisi ComboBox Provinsi"""
        try:
            self.chart_data = pd.read_excel(self.excel_file_path, sheet_name=0)
            print("Data dari file Excel berhasil dimuat.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal memuat file Excel untuk chart: {e}")
    
    def reload_excel_data(self):
        """Reload data ketika file Excel berubah."""
        try:
            print("Perubahan pada file Excel terdeteksi. Memuat ulang data...")
            
            # Tambahkan delay kecil untuk memastikan file tidak sedang ditulis ulang
            QTimer.singleShot(500, self.load_chart_data)
            
            # Update tampilan dengan data baru
            self.update_chart()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal memperbarui data: {e}")

    def update_stack_widget(self):
        """Mengganti tampilan QStackedWidget sesuai pilihan ComboBox"""
        selection = self.comboBox.currentText()
        
        # Menyesuaikan tampilan QStackedWidget berdasarkan pilihan ComboBox
        if selection == "Region":
            self.active_filter = "reg"
            self.stackedWidget.setCurrentWidget(self.data_pada_provinsi)
            self.stackedWidget_2.setCurrentWidget(self.keterangan_provinsi)
            self.stackedWidget_4.setCurrentWidget(self.dropdown_provinsi_3)
        elif selection == "Wilayah":
            self.active_filter = "wil"
            self.stackedWidget.setCurrentWidget(self.data_pada_wilayah)
            self.stackedWidget_2.setCurrentWidget(self.keterangan_wilayah)
            self.stackedWidget_4.setCurrentWidget(self.dropdown_wilayah_3)
        elif selection == "Estate":
            self.active_filter = "est"
            self.stackedWidget.setCurrentWidget(self.data_pada_estate)
            self.stackedWidget_2.setCurrentWidget(self.keterangan_estate)
            self.stackedWidget_4.setCurrentWidget(self.dropdown_estate_3)
        elif selection == "Divisi":
            self.active_filter = "div"
            self.stackedWidget.setCurrentWidget(self.data_pada_divisi)
            self.stackedWidget_2.setCurrentWidget(self.keterangan_divisi)
            self.stackedWidget_4.setCurrentWidget(self.dropdown_divisi_3)            
        
        # Memperbarui chart berdasarkan data yang difilter
        self.update_chart()

    def update_chart(self):
        
        # Tampilkan pemberitahuan hanya jika belum pernah ditampilkan
        if not self.notification_shown:
            QMessageBox.information(self, "Pemberitahuan", "Silahkan Pilih Filter dan Lokasi")
            self.notification_shown = True  # Set flag menjadi True

        """Update bar chart berdasarkan data dari file Excel dan pilihan region, wilayah, estate, atau divisi"""
        try:
            # Ambil pilihan utama dari comboBox untuk menentukan filter level
            selection = self.comboBox.currentText()

            # Konfigurasi berdasarkan pilihan level
            if selection == "Region":
                selected_value = self.tentukan_reg.text().strip()
                column_name = 'REGION'
                ulasan_label_name = "ulasan_reg"
                kesimpulan_label_name = 'kesimpulan_reg'
                frame = self.frame_18
                tahun_rotasi_input = self.tahun_reg.text().strip()
                hapus_input = self.hapus_reg.text().strip()  # Input hapus filter


            elif selection == "Wilayah":
                selected_value = self.tentukan_wil.text().strip()
                column_name = 'WILAYAH'
                ulasan_label_name = "ulasan_wil"
                kesimpulan_label_name = 'kesimpulan_wil'
                frame = self.frame_19
                tahun_rotasi_input = self.tahun_wil.text().strip()
                hapus_input = self.hapus_wil.text().strip()  # Input hapus filter


            elif selection == "Estate":
                selected_value = self.tentukan_est.text().strip()
                column_name = 'ESTATE'
                ulasan_label_name = "ulasan_est"
                kesimpulan_label_name = 'kesimpulan_est'
                frame = self.frame_23
                tahun_rotasi_input = self.tahun_est.text().strip()
                hapus_input = self.hapus_est.text().strip()  # Input hapus filter

            elif selection == "Divisi":
                selected_value = self.tentukan_div.text()
                column_name = 'DIVISI'
                ulasan_label_name = "ulasan_div"
                kesimpulan_label_name = 'kesimpulan_div'
                frame = self.frame_21
                tahun_rotasi_input = self.tahun_div.text().strip()
                hapus_input = self.hapus_div.text().strip()  # Input hapus filter
                
            else:
                return  # Jika tidak ada pilihan level yang sesuai, keluar dari fungsi

           # Validasi input lokasi
            if not selected_value:  # Jika input lokasi kosong
                QMessageBox.warning(self, "Peringatan", f"Silakan ketik lokasi pada kolom {selection}.")
                return

            # Filter data berdasarkan lokasi
            if column_name == "DIVISI":
                try:
                    selected_value = int(selected_value)  # Divisi diharapkan angka
                    filtered_data = self.chart_data[self.chart_data[column_name] == selected_value]
                except ValueError:
                    QMessageBox.warning(self, "Peringatan", f"Masukkan Divisi yang valid (angka).")
                    return
            else:
                filtered_data = self.chart_data[self.chart_data[column_name].str.contains(selected_value, case=False, na=False)]

            # Hapus data berdasarkan input hapus filter
            if hapus_input:
                hapus_filters = [item.strip() for item in hapus_input.split(',')]  # Pisahkan input berdasarkan koma
                for hapus_item in hapus_filters:
                    hapus_tahun_rotasi_match = re.match(r"^\d{4}(?:\s[Rr]\d)?$", hapus_item)  # Format valid: 2024, 2024 R1
                    if hapus_tahun_rotasi_match:
                        split_input = hapus_item.upper().split(" ")  # Konversi ke huruf besar
                        hapus_tahun = int(split_input[0]) if split_input[0].isdigit() else None
                        hapus_rotasi = split_input[1] if len(split_input) > 1 else None

                        # Hapus data berdasarkan Tahun dan Rotasi
                        if hapus_tahun and hapus_rotasi:
                            filtered_data = filtered_data[
                                ~((filtered_data['TAHUN'] == hapus_tahun) & (filtered_data['ROTASI'] == hapus_rotasi))
                            ]
                        elif hapus_tahun:
                            filtered_data = filtered_data[filtered_data['TAHUN'] != hapus_tahun]
                    else:
                        QMessageBox.warning(self, "Peringatan", f"Format tidak valid: {hapus_item}. Gunakan format 2024, 2024 R1, dll.")
                        return

            # Filter berdasarkan input tahun dan rotasi
            if tahun_rotasi_input:
                # Validasi input tahun dan rotasi
                tahun_rotasi_match = re.match(r"^\d{4}(?:\s[Rr]\d)?$", tahun_rotasi_input)  # Format valid: 2024, 2024 R1, r1
                rotasi_only_match = re.match(r"^[Rr]\d$", tahun_rotasi_input)  # Format valid: R1, r1

                if tahun_rotasi_match:
                    # Input berupa Tahun atau Tahun + Rotasi
                    split_input = tahun_rotasi_input.upper().split(" ")  # Konversi ke huruf besar
                    filter_tahun = int(split_input[0]) if split_input[0].isdigit() else None
                    filter_rotasi = split_input[1] if len(split_input) > 1 else None

                    # Filter berdasarkan Tahun jika ada
                    if filter_tahun:
                        filtered_data = filtered_data[filtered_data['TAHUN'] == filter_tahun]

                    # Filter berdasarkan Rotasi jika ada
                    if filter_rotasi:
                        filtered_data = filtered_data[filtered_data['ROTASI'] == filter_rotasi]

                elif rotasi_only_match:
                    # Input berupa Rotasi saja
                    filter_rotasi = tahun_rotasi_input.upper()  # Konversi ke huruf besar
                    filtered_data = filtered_data[filtered_data['ROTASI'] == filter_rotasi]

                else:
                    QMessageBox.warning(self, "Peringatan", "Masukkan format tahun/rotasi yang valid (contoh: 2024, 2024 R1, atau R1).")
                    return

            
            if filtered_data.empty:
                QMessageBox.warning(self, "Data Kosong", f"Tidak ada data yang cocok untuk {column_name}: {selected_value}")
                return

            # Mengelompokkan data berdasarkan Tahun dan Tingkat Kesehatan           
            grouped_data = (
                filtered_data.groupby(['TAHUN', 'ROTASI', 'KESEHATAN'])['Ha']
                .sum()
                .unstack(fill_value=0)
            )
                        
            # Sort index untuk urutan dinamis (tahun dan rotasi)
            grouped_data = grouped_data.sort_index(level=['TAHUN', 'ROTASI'], ascending=[True, True])

            # Menghitung total luas area berdasarkan tingkat kesehatan
            total_area = filtered_data['Ha'].sum()
            # area_by_health = filtered_data.groupby('KESEHATAN')['Ha'].sum()

            total_area_by_year_rotation = (
                filtered_data.groupby(['TAHUN', 'ROTASI'])['Ha']
                .sum()
                .sort_index(ascending=True)  # Pastikan urutannya sesuai tahun dan rotasi
            )
            
            ulasan = "Total luas area:"
            details_total = []
            for (tahun, rotasi), value in total_area_by_year_rotation.items():
                details_total.append(f"{tahun} {rotasi} ({value:.2f} Ha)")
            ulasan += ", ".join(details_total) + "\n"

            for health_level in grouped_data.columns:
                ulasan += f"- Luas area {health_level}: "
                details = []
                for (rotasi, tahun), value in grouped_data[health_level].items():
                    details.append(f"{rotasi} {tahun} ({value:.2f} Ha)")
                ulasan += ", ".join(details) + "\n"

            # Perbarui QLabel sesuai dengan widget aktif
            ulasan_label = self.findChild(QtWidgets.QLabel, ulasan_label_name)
            if ulasan_label:
                ulasan_label.setText(ulasan)
                ulasan_label.setStyleSheet("font-size: 8.5pt; font-family: Arial;")  # Atur font
                
            # KESIMPULAN 
            # KESIMPULAN DINAMIS
            kesimpulan = ""

            # Filter data tahun terbaru
            latest_year = grouped_data.index.get_level_values('TAHUN').max()
            latest_data = grouped_data.loc[latest_year]

            # Inisialisasi perhitungan total area per tingkat kesehatan
            total_area_per_health = grouped_data.sum()

            # Ambil tingkat kesehatan dengan area terbesar
            max_health_level = total_area_per_health.idxmax()

            # Tentukan kondisi berdasarkan tingkat kesehatan dominan
            if max_health_level == "Green":
                kesimpulan += "Kondisi lahan secara keseluruhan sangat baik dengan mayoritas area dalam kondisi Green.\n"
            elif max_health_level == "Moderate Green":
                kesimpulan += "Kondisi lahan secara keseluruhan cukup baik, meskipun sebagian besar area berada pada kondisi Moderate Green.\n"
            elif max_health_level == "Need Improvement":
                kesimpulan += "Kondisi lahan perlu perhatian karena mayoritas area berada dalam kondisi Need Improvement.\n"
            elif max_health_level == "Need Improvement Soon":
                kesimpulan += "Kondisi lahan perlu segera ditingkatkan karena sebagian besar area berada dalam kondisi Need Improvement Soon.\n"
            else:
                kesimpulan += "Kondisi lahan sulit ditentukan karena data tingkat kesehatan tidak merata.\n"

            #tambahkan analisis per rotasi
            for health_level in grouped_data.columns:
                kesimpulan += f"- Area {health_level}: "
                rotasi_values = grouped_data[health_level].dropna()  # Ambil data non-null
                
                if len(rotasi_values) > 1:  # Perlu minimal dua rotasi untuk membandingkan
                    # Urutkan rotasi berdasarkan tahun dan rotasi
                    sorted_rotasi_values = rotasi_values.sort_index(ascending=True)
                    
                    # Ambil rotasi terbaru dan sebelumnya
                    latest_rotasi = sorted_rotasi_values.index[-1]
                    previous_rotasi = sorted_rotasi_values.index[-2]
                    
                    latest_value = sorted_rotasi_values[latest_rotasi]
                    previous_value = sorted_rotasi_values[previous_rotasi]
                
                    # Hitung perbandingan
                    selisih = latest_value - previous_value
                    status = "kenaikan" if selisih > 0 else "penurunan"
                    kesimpulan += f"{status} {abs(selisih):.2f} Ha dibandingkan {previous_rotasi[0]} {previous_rotasi[1]}"
                else:
                    kesimpulan += "Tidak ada data sebelumnya untuk perbandingan."
                kesimpulan += "\n"

            # Perbarui QLabel untuk kesimpulan
            kesimpulan_label = self.findChild(QtWidgets.QLabel, kesimpulan_label_name)
            if kesimpulan_label:
                kesimpulan_label.setText(kesimpulan)
                kesimpulan_label.setStyleSheet("font-size: 8.5pt; font-family: Arial;")  # Atur font

            # Membuat canvas baru untuk bar chart
            self.canvas = MplCanvas(self, width=16, height=12, dpi=100)
            self.canvas.axes.clear()

            # Warna untuk tiap kategori tingkat kesehatan
            colors = ['#808080', '#1A9641', '#A6D96A', '#FDAE61', '#D7191C']
            bar_width = 0.175
            total_width = bar_width * len(grouped_data.columns)
            total_per_year = grouped_data.sum(axis=1)

            # Membuat bar chart untuk tiap tingkat kesehatan
            for i, (kesehatan, color) in enumerate(zip(grouped_data.columns, colors)):
                x_positions = [j + i * bar_width for j in range(len(grouped_data.index))]
                self.canvas.axes.bar(x_positions, grouped_data[kesehatan], color=color, width=bar_width, label=kesehatan)
                
                for x, y, total in zip(x_positions, grouped_data[kesehatan], total_per_year):
                    percentage = f"{(y / total * 100):.1f}%" if total > 0 else "0%"
                    self.canvas.axes.text(x, y + 0.02 * max(total_per_year), f"{y:.0f} Ha\n{percentage}", ha='center', va='bottom', fontsize=12, weight='bold', fontname='Arial')

        
            # Konfigurasi sumbu x berdasarkan kombinasi Rotasi dan Tahun
            x_ticks = [
                f"{tahun} {rotasi}"
                for tahun, rotasi in grouped_data.index
            ]
            x_positions = [
                j + total_width / 2 - bar_width / 2
                for j in range(len(grouped_data.index))
            ]
            self.canvas.axes.set_xticks(x_positions)
            self.canvas.axes.set_xticklabels(x_ticks, weight='bold', fontsize=14)
            
            # Konfigurasi tampilan chart
            self.canvas.axes.set_title(f"Grafik Kesehatan Tanaman per Rotasi dan Tahun - {selection} {selected_value}", weight='bold', fontsize=16)

            # Menambahkan padding agar teks tidak keluar dari grafik
            self.canvas.axes.margins(y=0.3)  # Tambah margin vertikal
            
            # Menempatkan legenda di bawah grafik
            self.canvas.axes.legend(loc='upper center', bbox_to_anchor=(0.5, -0.1), ncol=len(colors), frameon=False, fontsize=14)
            
            # Mengatur padding layout agar grafik dan teks berada di dalam area plot
            self.canvas.figure.tight_layout(pad=3)

            self.canvas.draw()

            # Menampilkan chart di frame yang diinginkan
            if frame.layout() is None:
                frame.setLayout(QtWidgets.QVBoxLayout())

            # Hapus widget lama di dalam frame
            for i in reversed(range(frame.layout().count())):
                frame.layout().itemAt(i).widget().setParent(None)

            # Tambahkan canvas ke dalam layout frame
            frame.layout().addWidget(self.canvas)
            
            self.is_analysis_done = True  # Set flag analisis ke True

        except Exception as e:
            QMessageBox.information(self, "Pemberitahuan", f"ERROR")

    def masukkan_data(self):
        """Memasukkan file (Excel, SHP, DBF), menampilkan di QTableWidget, dan konversi SHP/DBF ke CSV."""
        try:
            file_path, _ = QFileDialog.getOpenFileName(
                self, "Pilih File", "", "Supported Files (*.xlsx *.shp *.dbf)"
            )
            if not file_path:
                return

            file_extension = os.path.splitext(file_path)[1].lower()

            if file_extension in [".shp", ".dbf"]:
                self.convert_to_csv(file_path)  # Konversi SHP/DBF ke CSV
            elif file_extension == ".xlsx":
                self.load_excel(file_path)  # Tampilkan file Excel langsung
            else:
                QMessageBox.warning(self, "Error", "Format file tidak didukung.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal memuat file: {e}")
            print(f"Kesalahan: {e}")

    def convert_to_csv(self, file_path):
        """Konversi file SHP/DBF ke CSV dan tampilkan di QTableWidget dengan progress bar."""
        try:
            # Inisialisasi progress bar
            self.progress_load.setMinimum(0)
            self.progress_load.setMaximum(100)
            self.progress_load.setValue(0)
            self.progress_load.setVisible(True)

            # Langkah 1: Membaca file SHP atau DBF
            self.progress_load.setValue(10)
            gdf = gpd.read_file(file_path)
            self.progress_load.setValue(40)

            # Langkah 2: Buang kolom 'geometry' jika ada
            if 'geometry' in gdf.columns:
                gdf = gdf.drop(columns='geometry')

            # Langkah 3: Simpan sebagai CSV sementara
            self.progress_load.setValue(60)
            self.temp_csv_path = os.path.join(os.getcwd(), "temp_data.csv")
            gdf.to_csv(self.temp_csv_path, index=False)

            # Langkah 4: Tampilkan CSV di QTableWidget
            self.progress_load.setValue(80)
            self.load_csv(self.temp_csv_path)

            # Proses selesai
            self.progress_load.setValue(100)
            QMessageBox.information(self, "Sukses", "File berhasil dikonversi ke CSV dan ditampilkan.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal mengonversi file ke CSV: {e}")
        finally:
            # Setelah selesai, sembunyikan progress bar dengan sedikit delay agar terlihat rampung
            QtCore.QTimer.singleShot(500, lambda: self.progress_load.setVisible(False))
            self.progress_load.setValue(0)

    def load_csv(self, file_path):
        """Memuat data CSV ke QTableWidget."""
        try:
            df = pd.read_csv(file_path)

            # Tampilkan di QTableWidget
            self.display_in_table(df)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal memuat file CSV: {e}")

    def rekap_data(self):
        """Merekap data untuk menghitung total PKK, luas tanam (Ha), dan persentase (Persen) berdasarkan tingkat kesehatan untuk setiap BLOK_SAP."""
        try:
            # Pastikan file areal statement sudah dimuat
            if not hasattr(self, "areal_statement") or self.areal_statement is None:
                QMessageBox.warning(self, "Error", "File areal statement belum dimuat. Silakan pilih file terlebih dahulu.")
                return

            # Validasi kolom yang dibutuhkan dalam areal statement
            required_areal_columns = ['REGION', 'WILAYAH', 'BLOK_ID', 'TANAM', 'BLOK_SAP', 'LUAS_TANAM', 'ESTATE', 'DIVISI']
            missing_areal_columns = [col for col in required_areal_columns if col not in self.areal_statement.columns]
            if missing_areal_columns:
                QMessageBox.warning(self, "Error", f"Kolom berikut tidak ditemukan di file areal statement: {', '.join(missing_areal_columns)}")
                return

            # Pastikan temp_csv_path ada dan file tersedia
            if not hasattr(self, "temp_csv_path"):
                QMessageBox.warning(self, "Peringatan", "Path file tidak ditemukan.")
                return

            if not os.path.exists(self.temp_csv_path):
                QMessageBox.warning(self, "Peringatan", "File CSV sementara tidak ditemukan.")
                return

            # Baca data dari file CSV yang sudah dikonversi
            df = pd.read_csv(self.temp_csv_path)
            print(f"Data yang dibaca: {df.head()}")  # Debugging

            # Validasi kolom yang dibutuhkan dalam data CSV
            required_columns = ['ESTATE', 'DIVISI', 'BLOK_SAP', 'KESEHATAN', 'ROTASI', 'TAHUN']
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                QMessageBox.warning(self, "Error", f"Kolom berikut tidak ditemukan di data: {', '.join(missing_columns)}")
                return

            # **Mengubah kolom ROTASI menjadi format string R1, R2, dst.**
            df['ROTASI'] = df['ROTASI'].apply(lambda x: f"R{x}")

            # Hitung jumlah baris dengan KESEHATAN yang sama untuk setiap BLOK_SAP
            rekap_data = (
                df.groupby(['ESTATE', 'DIVISI', 'BLOK_SAP', 'KESEHATAN', 'ROTASI', 'TAHUN'], as_index=False)
                .size()  # Menghitung jumlah baris
                .rename(columns={'size': 'PKK'})  # Mengganti nama kolom size menjadi PKK
                .sort_values(by=['ESTATE', 'DIVISI', 'BLOK_SAP', 'KESEHATAN'])  # Urutkan data
            )

            # Hitung total PKK per kombinasi ESTATE, DIVISI, dan BLOK_SAP
            rekap_data['Total_PKK'] = rekap_data.groupby(['ESTATE', 'DIVISI', 'BLOK_SAP'])['PKK'].transform('sum')

            # Gabungkan dengan file areal statement
            merged_data = pd.merge(
                rekap_data,
                self.areal_statement[['REGION', 'WILAYAH', 'BLOK_ID', 'TANAM','ESTATE', 'DIVISI', 'BLOK_SAP', 'LUAS_TANAM']],  # Ambil kolom yang relevan
                on=['ESTATE', 'DIVISI', 'BLOK_SAP'],
                how='left'
            )

            # Validasi apakah ada blok yang tidak ditemukan di areal statement
            if merged_data['LUAS_TANAM'].isnull().any():
                missing_blocks = merged_data[merged_data['LUAS_TANAM'].isnull()]['BLOK_SAP'].unique()
                QMessageBox.warning(self, "Warning", f"Beberapa BLOK_SAP tidak ditemukan di file areal statement: {', '.join(missing_blocks)}")
                merged_data = merged_data.dropna(subset=['LUAS_TANAM'])  # Hapus baris dengan LUAS_TANAM NaN

            # Hitung luas tanam (Ha) berdasarkan proporsi PKK
            merged_data['Ha'] = merged_data['LUAS_TANAM'] * (merged_data['PKK'] / merged_data['Total_PKK'])

            # Format kolom Ha menjadi 2 angka di belakang koma
            merged_data['Ha'] = merged_data['Ha'].round(14)

            # Hitung total Ha per kombinasi ESTATE, DIVISI, dan BLOK_SAP
            merged_data['Total_Ha'] = merged_data.groupby(['ESTATE', 'DIVISI', 'BLOK_SAP'])['Ha'].transform('sum')

            # Hitung persentase
            merged_data['Persen'] = (merged_data['Ha'] / merged_data['Total_Ha']).round(15)

            # Pilih hanya kolom yang akan ditampilkan
            final_data = merged_data[['REGION', 'WILAYAH', 'BLOK_ID', 'TANAM', 'ESTATE', 'DIVISI', 'BLOK_SAP', 'KESEHATAN', 'ROTASI', 'TAHUN', 'PKK', 'Ha', 'Persen']]

            # Tampilkan data baru di tableWidget_rekap
            self.display_in_tablewidget_rekap(final_data)

            QMessageBox.information(self, "Sukses", "Data rekap berhasil dibuat dengan kolom PKK, Ha, dan Persen ditambahkan.")
            print("Data rekap berhasil dibuat:", final_data.head())

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal membuat rekap data: {e}")
            print(f"Kesalahan saat rekap data: {e}")


    def display_in_table(self, dataframe):
        """Tampilkan DataFrame di QTableWidget."""
        try:
            self.tableWidget.clear()
            self.tableWidget.setRowCount(dataframe.shape[0])
            self.tableWidget.setColumnCount(dataframe.shape[1])
            self.tableWidget.setHorizontalHeaderLabels(dataframe.columns)

            for row_idx, row in dataframe.iterrows():
                for col_idx, value in enumerate(row):
                    self.tableWidget.setItem(row_idx, col_idx, QTableWidgetItem(str(value)))

            self.tableWidget.resizeColumnsToContents()
            self.tableWidget.resizeRowsToContents()
            print("Data berhasil ditampilkan di QTableWidget.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal menampilkan data: {e}")
            print(f"Kesalahan saat menampilkan data: {e}")

    def display_in_tablewidget_rekap(self, dataframe):
        """Tampilkan DataFrame di tableView (QTableView) di dalam tableWidget_rekap."""
        try:
            if self.tableView_rekap is None:
                self.tableView_rekap = self.tableWidget_rekap.findChild(QTableView, "tableView")

            if self.tableView_rekap is None:
                QMessageBox.critical(self, "Error", "tableView tidak ditemukan di tableWidget_rekap.")
                return

            # Set model DataFrame untuk tableView
            model = PandasModel(dataframe)
            self.tableView_rekap.setModel(model)
            self.tableView_rekap.resizeColumnsToContents()
            self.tableView_rekap.resizeRowsToContents()

            print("Data rekap berhasil ditampilkan di tableView.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal menampilkan data: {e}")
            print(f"Kesalahan saat menampilkan data: {e}")            
    
    def simpan_data(self):
        """Simpan data hasil rekap ke file master Excel dan muat ulang data tanpa mengubah urutan kolom."""
        try:
            if self.tableView_rekap is None or self.tableView_rekap.model() is None:
                QMessageBox.warning(self, "Peringatan", "Tidak ada data rekap untuk disimpan!")
                return

            # Konversi model data di tableView_rekap ke DataFrame
            model = self.tableView_rekap.model()
            rows = model.rowCount()
            cols = model.columnCount()
            rekap_data = pd.DataFrame(
                [[model.data(model.index(row, col)) for col in range(cols)] for row in range(rows)],
                columns=[model.headerData(col, Qt.Horizontal) for col in range(cols)]
            )

            if rekap_data.empty:
                QMessageBox.warning(self, "Peringatan", "Data rekap kosong!")
                return

            # Path file master Excel
            file_master_path = self.excel_file_path
            sheet_name = "Sheet1"

            # Baca file master Excel jika ada
            if os.path.exists(file_master_path):
                existing_data = pd.read_excel(file_master_path, sheet_name=sheet_name)
            else:
                # Jika file master belum ada, buat DataFrame kosong dengan kolom dari data rekap
                existing_data = pd.DataFrame(columns=rekap_data.columns)

            # Simpan urutan asli kolom dari file master
            original_columns = list(existing_data.columns)

            # Tambahkan kolom yang hilang di file maste
            # r (jika ada kolom baru di rekap_data)
            for col in rekap_data.columns:
                if col not in existing_data.columns:
                    existing_data[col] = None

            # Pastikan semua kolom dari existing_data tetap dalam urutan aslinya
            combined_data = pd.concat([existing_data, rekap_data], ignore_index=True)

            # Hapus duplikasi berdasarkan semua kolom
            combined_data = combined_data.drop_duplicates(keep=False)

            # Atur ulang kolom agar sesuai dengan urutan asli, tambahkan kolom baru di akhir
            final_columns = original_columns + [col for col in combined_data.columns if col not in original_columns]
            combined_data = combined_data[final_columns]

            # Simpan data ke file master Excel
            with pd.ExcelWriter(file_master_path, engine='openpyxl') as writer:
                combined_data.to_excel(writer, index=False, sheet_name=sheet_name)

            # Muat ulang data ke dalam aplikasi
            self.chart_data = pd.read_excel(file_master_path, sheet_name=sheet_name)

            QMessageBox.information(self, "Sukses", "Data berhasil disimpan tanpa duplikasi dan dengan urutan kolom tetap.")
            print(f"Data terbaru berhasil dimuat: {self.chart_data.head()}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal menyimpan atau memperbarui data: {e}")
            print(f"Kesalahan saat menyimpan data: {e}")
    


    def download(self):
        """Mengunduh hasil analisis berdasarkan filter aktif."""
        try:
            # Validasi apakah analisis sudah dilakukan
            if not self.is_analysis_done:
                QMessageBox.warning(self, "Peringatan", "Silakan lakukan analisis terlebih dahulu sebelum mengunduh hasil!")
                return

            # Tentukan ulasan dan kesimpulan berdasarkan filter aktif
            ulasan_label_name = f"ulasan_{self.active_filter}"  # Contoh: ulasan_reg, ulasan_wil, dll.
            kesimpulan_label_name = f"kesimpulan_{self.active_filter}"  # Contoh: kesimpulan_reg, kesimpulan_wil, dll.

            # Ambil QLabel untuk ulasan dan kesimpulan
            ulasan_label = self.findChild(QtWidgets.QLabel, ulasan_label_name)
            kesimpulan_label = self.findChild(QtWidgets.QLabel, kesimpulan_label_name)

            # Validasi keberadaan QLabel
            if not ulasan_label or not kesimpulan_label:
                QMessageBox.warning(self, "Error", f"Ulasan atau kesimpulan untuk {self.active_filter} tidak tersedia!")
                return

            # Ambil teks ulasan dan kesimpulan
            ulasan_text = ulasan_label.text()
            kesimpulan_text = kesimpulan_label.text()

            # Buat dokumen Word baru
            doc = Document()
            doc.add_heading("Hasil Analisis", level=1)

            # Tambahkan bar chart (jika tersedia)
            if hasattr(self, 'canvas') and self.canvas:
                image_buffer = io.BytesIO()
                self.canvas.figure.savefig(image_buffer, format='png')
                image_buffer.seek(0)
                doc.add_heading("Bar Chart", level=2)
                doc.add_picture(image_buffer, width=Inches(6))
                image_buffer.close()
            else:
                QMessageBox.warning(self, "Peringatan", "Bar chart tidak tersedia.")

            # Tambahkan ulasan
            doc.add_heading("Ulasan", level=2)
            doc.add_paragraph(ulasan_text)

            # Tambahkan kesimpulan
            doc.add_heading("Kesimpulan", level=2)
            doc.add_paragraph(kesimpulan_text)

            # Pilih lokasi penyimpanan
            save_path, _ = QFileDialog.getSaveFileName(self, "Simpan Hasil Analisis", "", "Word Documents (*.docx)")
            if save_path:
                if not save_path.endswith(".docx"):
                    save_path += ".docx"
                doc.save(save_path)
                QMessageBox.information(self, "Sukses", f"Hasil analisis berhasil disimpan di: {save_path}")
            else:
                QMessageBox.warning(self, "Batal", "Proses penyimpanan dibatalkan.")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Gagal mengunduh hasil analisis: {e}")
            print(f"Kesalahan saat mengunduh hasil analisis: {e}")


    def toggle_menu(self):
        current_visibility = self.left_menu_widget.isVisible()
        self.left_menu_widget.setVisible(not current_visibility)

    def minimize_window(self):
        self.showMinimized()
    
    def toggle_fullscreen(self):
        """Toggle fullscreen mode for the application window."""
        if self.is_fullscreen:
            self.showNormal()
            self.is_fullscreen = False
        else:
            self.showMaximized()
            self.is_fullscreen = True

    def close_window(self):
        """Close the application window."""
        self.close()

    def mousePressEvent(self, event):
        """Enable drag for the frameless window on mouse press event."""
        if not self.is_fullscreen and event.button() == QtCore.Qt.LeftButton:
            self.is_dragging = True
            self.drag_position = event.globalPos() - self.frameGeometry().topLeft()
            event.accept()

    def mouseMoveEvent(self, event):
        """Handle mouse move event to drag the window if frameless."""
        if not self.is_fullscreen and self.is_dragging:
            self.move(event.globalPos() - self.drag_position)
            event.accept()

    def mouseReleaseEvent(self, event):
        """Disable dragging on mouse release."""
        self.is_dragging = False


# Jalankan aplikasi
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainApp()
    window.show()
    window.showMaximized()
    sys.exit(app.exec_())